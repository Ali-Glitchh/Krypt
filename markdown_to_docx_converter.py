"""
Markdown to DOCX Converter for KoinToss FYP Report
Converts the Markdown FYP report to a properly formatted Word document.
"""

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from bs4 import BeautifulSoup
import re
import os

def create_styles(doc):
    """Create custom styles for the document"""
    styles = doc.styles
    
    # Title style
    try:
        title_style = styles['Title']
    except KeyError:
        title_style = styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Heading 1 style
    try:
        h1_style = styles['Heading 1']
    except KeyError:
        h1_style = styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.paragraph_format.space_before = Pt(12)
    h1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style
    try:
        h2_style = styles['Heading 2']
    except KeyError:
        h2_style = styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.paragraph_format.space_before = Pt(6)
    h2_style.paragraph_format.space_after = Pt(3)
    
    # Normal style
    try:
        normal_style = styles['Normal']
    except KeyError:
        normal_style = styles.add_style('Normal', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.line_spacing = 1.15

def process_markdown_content(markdown_content):
    """Process markdown content and convert to HTML"""
    # Convert markdown to HTML
    html = markdown.markdown(markdown_content, extensions=['tables', 'fenced_code'])
    
    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html, 'html.parser')
    return soup

def add_element_to_doc(doc, element):
    """Add HTML element to Word document"""
    if element.name == 'h1':
        p = doc.add_paragraph(element.get_text(), style='Heading 1')
    elif element.name == 'h2':
        p = doc.add_paragraph(element.get_text(), style='Heading 2')
    elif element.name == 'h3':
        p = doc.add_paragraph(element.get_text(), style='Heading 2')
    elif element.name == 'p':
        text = element.get_text()
        if text.strip():  # Only add non-empty paragraphs
            p = doc.add_paragraph(text, style='Normal')
    elif element.name == 'ul':
        for li in element.find_all('li'):
            p = doc.add_paragraph(li.get_text(), style='List Bullet')
    elif element.name == 'ol':
        for i, li in enumerate(element.find_all('li'), 1):
            p = doc.add_paragraph(f"{i}. {li.get_text()}", style='List Number')
    elif element.name == 'pre':
        # Handle code blocks
        code_text = element.get_text()
        p = doc.add_paragraph(code_text)
        p.style.font.name = 'Courier New'
        p.style.font.size = Pt(10)
    elif element.name == 'table':
        # Handle tables
        rows = element.find_all('tr')
        if rows:
            table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['th', 'td'])))
            table.style = 'Table Grid'
            
            for i, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for j, cell in enumerate(cells):
                    if i < len(table.rows) and j < len(table.columns):
                        table.cell(i, j).text = cell.get_text().strip()

def convert_markdown_to_docx(input_file, output_file):
    """Convert Markdown file to DOCX"""
    print(f"Converting {input_file} to {output_file}...")
    
    # Read markdown file
    with open(input_file, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # Create new document
    doc = Document()
    
    # Create styles
    create_styles(doc)
    
    # Process markdown content
    soup = process_markdown_content(markdown_content)
    
    # Add elements to document
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'pre', 'table']):
        add_element_to_doc(doc, element)
    
    # Handle Mermaid diagrams - convert to text blocks for now
    mermaid_pattern = r'```mermaid\n(.*?)\n```'
    mermaid_matches = re.findall(mermaid_pattern, markdown_content, re.DOTALL)
    
    if mermaid_matches:
        print(f"Found {len(mermaid_matches)} Mermaid diagrams")
        # Add a note about diagrams
        p = doc.add_paragraph("\n[Note: The following Mermaid diagrams were included in the original document. They should be converted to images for the final submission.]")
        p.style.font.italic = True
        
        for i, diagram in enumerate(mermaid_matches, 1):
            p = doc.add_paragraph(f"\nDiagram {i}:")
            p.style.font.bold = True
            p = doc.add_paragraph(diagram)
            p.style.font.name = 'Courier New'
            p.style.font.size = Pt(9)
    
    # Save document
    doc.save(output_file)
    print(f"Successfully converted to {output_file}")

def main():
    """Main conversion function"""
    input_file = "KoinToss_FYP_Report.md"
    output_file = "KoinToss_FYP_Report_Converted.docx"
    
    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found!")
        return
    
    try:
        convert_markdown_to_docx(input_file, output_file)
        print(f"\nConversion completed successfully!")
        print(f"Output file: {output_file}")
        print("\nNote: Please review the converted document and:")
        print("1. Convert Mermaid diagrams to proper images")
        print("2. Adjust formatting as needed")
        print("3. Add any missing elements from the original template")
        
    except Exception as e:
        print(f"Error during conversion: {e}")

if __name__ == "__main__":
    main()
