"""
Enhanced Markdown to DOCX Converter for KoinToss FYP Report
Creates a professional FYP document following CS department standards.
"""

import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from bs4 import BeautifulSoup
import re
import os
from datetime import datetime

def setup_document_properties(doc):
    """Setup document properties and margins"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)

def create_professional_styles(doc):
    """Create professional styles matching FYP requirements"""
    styles = doc.styles
    
    # Title Page Title
    title_style = styles.add_style('FYP Title', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(18)
    
    # Chapter Heading
    chapter_style = styles.add_style('Chapter Heading', WD_STYLE_TYPE.PARAGRAPH)
    chapter_style.font.name = 'Times New Roman'
    chapter_style.font.size = Pt(16)
    chapter_style.font.bold = True
    chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    chapter_style.paragraph_format.space_before = Pt(24)
    chapter_style.paragraph_format.space_after = Pt(12)
    
    # Section Heading
    section_style = styles.add_style('Section Heading', WD_STYLE_TYPE.PARAGRAPH)
    section_style.font.name = 'Times New Roman'
    section_style.font.size = Pt(14)
    section_style.font.bold = True
    section_style.paragraph_format.space_before = Pt(12)
    section_style.paragraph_format.space_after = Pt(6)
    
    # Subsection Heading
    subsection_style = styles.add_style('Subsection Heading', WD_STYLE_TYPE.PARAGRAPH)
    subsection_style.font.name = 'Times New Roman'
    subsection_style.font.size = Pt(12)
    subsection_style.font.bold = True
    subsection_style.paragraph_format.space_before = Pt(6)
    subsection_style.paragraph_format.space_after = Pt(3)
    
    # Body Text
    body_style = styles.add_style('FYP Body', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = 'Times New Roman'
    body_style.font.size = Pt(12)
    body_style.paragraph_format.line_spacing = 1.5
    body_style.paragraph_format.space_after = Pt(6)
    body_style.paragraph_format.first_line_indent = Inches(0.5)
    
    # Code Style
    code_style = styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)

def add_title_page(doc):
    """Add a professional title page"""
    # University name
    p = doc.add_paragraph("UNIVERSITY OF TECHNOLOGY MALAYSIA", style='FYP Title')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Faculty
    p = doc.add_paragraph("FACULTY OF COMPUTING", style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.bold = True
    
    # Add some space
    doc.add_paragraph("\n\n")
    
    # Project title
    p = doc.add_paragraph("KOINTOSS: CRYPTOCURRENCY CHATBOT WITH ADVANCED TRADING INSIGHTS", style='FYP Title')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add space
    doc.add_paragraph("\n\n")
    
    # Project details
    p = doc.add_paragraph("FINAL YEAR PROJECT REPORT", style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.bold = True
    
    doc.add_paragraph("\n")
    
    # Student details (placeholder)
    p = doc.add_paragraph("Prepared by:", style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    
    p = doc.add_paragraph("[Student Name]", style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    
    p = doc.add_paragraph("[Student ID]", style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    
    doc.add_paragraph("\n")
    
    # Supervisor
    p = doc.add_paragraph("Supervised by:", style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    
    p = doc.add_paragraph("[Supervisor Name]", style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    
    doc.add_paragraph("\n\n")
    
    # Date
    p = doc.add_paragraph(f"{datetime.now().strftime('%B %Y')}", style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    
    # Page break
    doc.add_page_break()

def process_enhanced_markdown(markdown_content):
    """Enhanced markdown processing with better structure recognition"""
    # Convert markdown to HTML with extensions
    html = markdown.markdown(
        markdown_content, 
        extensions=['tables', 'fenced_code', 'toc', 'codehilite']
    )
    
    # Parse HTML
    soup = BeautifulSoup(html, 'html.parser')
    return soup

def add_enhanced_element(doc, element, chapter_counter=[0]):
    """Add elements with enhanced formatting"""
    if element.name == 'h1':
        chapter_counter[0] += 1
        text = element.get_text()
        p = doc.add_paragraph(f"CHAPTER {chapter_counter[0]}: {text.upper()}", style='Chapter Heading')
        
    elif element.name == 'h2':
        text = element.get_text()
        p = doc.add_paragraph(text, style='Section Heading')
        
    elif element.name == 'h3':
        text = element.get_text()
        p = doc.add_paragraph(text, style='Subsection Heading')
        
    elif element.name == 'p':
        text = element.get_text().strip()
        if text:
            p = doc.add_paragraph(text, style='FYP Body')
            
    elif element.name == 'ul':
        for li in element.find_all('li'):
            text = li.get_text().strip()
            if text:
                p = doc.add_paragraph(text, style='List Bullet')
                p.style.font.name = 'Times New Roman'
                p.style.font.size = Pt(12)
                
    elif element.name == 'ol':
        for i, li in enumerate(element.find_all('li'), 1):
            text = li.get_text().strip()
            if text:
                p = doc.add_paragraph(f"{i}. {text}", style='List Number')
                p.style.font.name = 'Times New Roman'
                p.style.font.size = Pt(12)
                
    elif element.name == 'pre':
        code_text = element.get_text()
        if code_text.strip():
            p = doc.add_paragraph(code_text, style='Code Block')
            
    elif element.name == 'table':
        rows = element.find_all('tr')
        if rows:
            # Determine number of columns from first row
            cols = len(rows[0].find_all(['th', 'td']))
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = 'Table Grid'
            
            # Set table formatting
            for row in table.rows:
                row.height = Inches(0.3)
                for cell in row.cells:
                    cell.vertical_alignment = 1  # Center alignment
                    
            # Fill table data
            for i, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for j, cell in enumerate(cells):
                    if j < len(table.columns):
                        table.cell(i, j).text = cell.get_text().strip()
                        # Make header row bold
                        if i == 0:
                            table.cell(i, j).paragraphs[0].runs[0].font.bold = True

def add_diagram_placeholders(doc, markdown_content):
    """Add placeholders for Mermaid diagrams"""
    mermaid_pattern = r'```mermaid\n(.*?)\n```'
    mermaid_matches = re.findall(mermaid_pattern, markdown_content, re.DOTALL)
    
    if mermaid_matches:
        # Add diagrams section
        p = doc.add_paragraph("SYSTEM DIAGRAMS", style='Section Heading')
        
        p = doc.add_paragraph(
            "The following diagrams illustrate the system architecture, data flow, and component interactions of the KoinToss cryptocurrency chatbot system.",
            style='FYP Body'
        )
        
        for i, diagram in enumerate(mermaid_matches, 1):
            # Determine diagram type from content
            diagram_type = "System Diagram"
            if "graph" in diagram.lower():
                diagram_type = "System Architecture Diagram"
            elif "sequenceDiagram" in diagram:
                diagram_type = "Sequence Diagram"
            elif "classDiagram" in diagram:
                diagram_type = "Class Diagram"
            elif "flowchart" in diagram.lower():
                diagram_type = "Flowchart"
                
            p = doc.add_paragraph(f"{diagram_type} {i}", style='Subsection Heading')
            
            # Add diagram placeholder
            p = doc.add_paragraph(f"[{diagram_type} {i} - To be inserted]", style='Normal')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
            # Add diagram code as reference
            p = doc.add_paragraph("Diagram Source Code:", style='Normal')
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(10)
            
            p = doc.add_paragraph(diagram, style='Code Block')
            
            doc.add_paragraph("")  # Add space

def convert_enhanced_markdown_to_docx(input_file, output_file):
    """Enhanced conversion with professional formatting"""
    print(f"Converting {input_file} to {output_file} with enhanced formatting...")
    
    # Read markdown file
    with open(input_file, 'r', encoding='utf-8') as f:
        markdown_content = f.read()
    
    # Create new document
    doc = Document()
    
    # Setup document
    setup_document_properties(doc)
    create_professional_styles(doc)
    
    # Add title page
    add_title_page(doc)
    
    # Process content
    soup = process_enhanced_markdown(markdown_content)
    
    # Add table of contents placeholder
    p = doc.add_paragraph("TABLE OF CONTENTS", style='Chapter Heading')
    p = doc.add_paragraph("[Table of Contents - To be generated in Word]", style='Normal')
    p.runs[0].font.italic = True
    doc.add_page_break()
    
    # Process main content
    chapter_counter = [0]
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'ul', 'ol', 'pre', 'table']):
        add_enhanced_element(doc, element, chapter_counter)
    
    # Add diagrams section
    add_diagram_placeholders(doc, markdown_content)
    
    # Save document
    doc.save(output_file)
    print(f"Successfully created enhanced FYP document: {output_file}")

def main():
    """Main function for enhanced conversion"""
    input_file = "KoinToss_FYP_Report.md"
    output_file = "KoinToss_FYP_Report_Professional.docx"
    
    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found!")
        return
    
    try:
        convert_enhanced_markdown_to_docx(input_file, output_file)
        print(f"\nEnhanced conversion completed successfully!")
        print(f"Output file: {output_file}")
        print("\nFinal steps for submission:")
        print("1. Open the document in Microsoft Word")
        print("2. Generate automatic Table of Contents")
        print("3. Convert Mermaid diagrams to images and insert them")
        print("4. Update student name, ID, and supervisor details on title page")
        print("5. Add page numbers")
        print("6. Review formatting and make final adjustments")
        print("7. Export as PDF for electronic submission")
        
    except Exception as e:
        print(f"Error during enhanced conversion: {e}")

if __name__ == "__main__":
    main()
